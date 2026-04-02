"""
Router Files - liste et rendu des fichiers de conversation.
"""
import base64
import io
import logging
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Header, HTTPException
from pydantic import BaseModel

from core.auth import get_user_id_from_header
from core.database import Conversation, ConversationFile, Message, SessionLocal
from core.firebase_client import firebase_storage as storage

router = APIRouter(prefix="/files", tags=["Files"])
logger = logging.getLogger(__name__)


class InlineFilePayload(BaseModel):
    data_url: str
    name: Optional[str] = None


def _utc_iso(dt: Optional[datetime]) -> Optional[str]:
    if dt is None:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    else:
        dt = dt.astimezone(timezone.utc)
    return dt.isoformat().replace("+00:00", "Z")


def _normalize_kind(kind: Optional[str], mime_type: Optional[str] = None, file_name: Optional[str] = None) -> str:
    raw_kind = str(kind or "").strip().lower()
    raw_mime = str(mime_type or "").strip().lower()
    lower_name = str(file_name or "").strip().lower()

    if raw_kind in {"image", "video", "audio", "document", "sheet", "spreadsheet"}:
        return "sheet" if raw_kind == "spreadsheet" else raw_kind
    if raw_mime.startswith("image/") or lower_name.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
        return "image"
    if raw_mime.startswith("video/") or lower_name.endswith((".mp4", ".mov", ".webm", ".mkv")):
        return "video"
    if raw_mime.startswith("audio/") or lower_name.endswith((".mp3", ".wav", ".ogg", ".m4a")):
        return "audio"
    if "spreadsheet" in raw_mime or "excel" in raw_mime or lower_name.endswith((".xlsx", ".xls", ".csv")):
        return "sheet"
    if "wordprocessingml" in raw_mime or raw_mime.endswith("/document") or lower_name.endswith((".docx", ".doc", ".pdf", ".txt", ".md")):
        return "document"
    return "document"


def _build_data_url(raw_b64: Optional[str], mime_type: Optional[str]) -> Optional[str]:
    if not raw_b64 or not mime_type:
        return None
    return f"data:{mime_type};base64,{raw_b64}"


def _estimate_size_from_base64(raw_b64: Optional[str]) -> Optional[int]:
    if not raw_b64:
        return None
    padding = raw_b64.count("=")
    return max(0, int(len(raw_b64) * 3 / 4) - padding)


def _serialize_conversation_file(row: ConversationFile, conv_titles: Dict[str, str]) -> Dict[str, Any]:
    kind = _normalize_kind(row.file_type, row.mime_type, row.file_name)
    return {
        "id": row.id,
        "name": row.file_name,
        "url": row.file_url,
        "remote_url": row.file_url,
        "data_url": None,
        "type": row.mime_type or row.file_type or kind,
        "kind": kind,
        "mime_type": row.mime_type,
        "created_at": _utc_iso(row.created_at),
        "size": row.file_size,
        "conversation_id": row.conversation_id,
        "conversation_title": conv_titles.get(row.conversation_id, ""),
        "source": "conversation_file",
        "inline": False,
        "ephemeral": False,
    }


def _serialize_message_attachment(row: Message, conv_titles: Dict[str, str]) -> Optional[Dict[str, Any]]:
    attachment = row.attachment_json
    if not isinstance(attachment, dict):
        return None

    mime_type = str(attachment.get("type") or attachment.get("mime_type") or "").strip()
    name = str(attachment.get("name") or "").strip() or f"media-{row.id}"
    kind = _normalize_kind(attachment.get("kind"), mime_type, name)
    raw_b64 = attachment.get("data")
    data_url = (
        attachment.get("data_url")
        or attachment.get("dataUrl")
        or _build_data_url(raw_b64, mime_type)
    )
    remote_url = attachment.get("url")
    effective_url = remote_url or data_url
    if not effective_url:
        return None

    return {
        "id": f"message-{row.id}",
        "name": name,
        "url": effective_url,
        "remote_url": remote_url,
        "data_url": data_url,
        "type": mime_type or kind,
        "kind": kind,
        "mime_type": mime_type or None,
        "created_at": _utc_iso(row.timestamp),
        "size": _estimate_size_from_base64(raw_b64),
        "conversation_id": row.conversation_id,
        "conversation_title": conv_titles.get(row.conversation_id, ""),
        "source": "message_attachment",
        "inline": not bool(remote_url),
        "ephemeral": bool(attachment.get("ephemeral")),
    }


def _merge_file_entries(entries: List[Optional[Dict[str, Any]]], limit: int = 200) -> List[Dict[str, Any]]:
    merged: Dict[str, Dict[str, Any]] = {}

    for entry in entries:
        if not entry:
            continue

        identity = (
            entry.get("remote_url")
            or entry.get("data_url")
            or entry.get("url")
            or f"{entry.get('conversation_id', '')}::{entry.get('name', '')}::{entry.get('created_at', '')}"
        )
        existing = merged.get(identity)
        if not existing:
            merged[identity] = entry
            continue

        if not existing.get("remote_url") and entry.get("remote_url"):
            existing["remote_url"] = entry.get("remote_url")
            existing["url"] = entry.get("remote_url") or existing.get("url")
            existing["inline"] = False
        if not existing.get("data_url") and entry.get("data_url"):
            existing["data_url"] = entry.get("data_url")
        if not existing.get("mime_type") and entry.get("mime_type"):
            existing["mime_type"] = entry.get("mime_type")
            existing["type"] = entry.get("type") or existing.get("type")
        if not existing.get("size") and entry.get("size"):
            existing["size"] = entry.get("size")
        if not existing.get("conversation_title") and entry.get("conversation_title"):
            existing["conversation_title"] = entry.get("conversation_title")
        if existing.get("source") != "conversation_file" and entry.get("source") == "conversation_file":
            existing["source"] = "conversation_file"

    return sorted(
        merged.values(),
        key=lambda item: item.get("created_at") or "",
        reverse=True,
    )[:limit]


async def _load_file_bytes(url: Optional[str] = None, data_url: Optional[str] = None) -> bytes:
    if data_url:
        try:
            _, raw_b64 = data_url.split(",", 1)
            return base64.b64decode(raw_b64)
        except Exception as exc:
            raise HTTPException(status_code=400, detail="Fichier inline invalide.") from exc

    if not url:
        raise HTTPException(status_code=400, detail="Aucune source de fichier fournie.")

    import httpx

    async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as http:
        resp = await http.get(url)
        if resp.status_code == 404:
            raise HTTPException(status_code=404, detail="Fichier introuvable (404).")
        resp.raise_for_status()
        return resp.content


def _render_docx_bytes_as_html(file_bytes: bytes) -> dict:
    import html as html_lib
    from docx import Document as DocxDoc

    doc = DocxDoc(io.BytesIO(file_bytes))
    html_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name
        if style == "Title":
            html_parts.append(f'<h1 style="font-size:2em;font-weight:bold;margin:0.5em 0">{html_lib.escape(text)}</h1>')
        elif style.startswith("Heading "):
            try:
                level = min(int(style.split()[-1]), 6)
            except ValueError:
                level = 2
            sizes = {1: "1.6em", 2: "1.35em", 3: "1.15em", 4: "1.05em", 5: "1em", 6: "0.9em"}
            html_parts.append(
                f'<h{level} style="font-size:{sizes.get(level, "1em")};font-weight:bold;margin:0.8em 0 0.3em">{html_lib.escape(text)}</h{level}>'
            )
        elif "Bullet" in style or "bullet" in style:
            html_parts.append(f'<li style="margin-left:1.5em;list-style-type:disc">{html_lib.escape(text)}</li>')
        elif "Number" in style or "number" in style:
            html_parts.append(f'<li style="margin-left:1.5em;list-style-type:decimal">{html_lib.escape(text)}</li>')
        else:
            parts_html = []
            for run in para.runs:
                if run.text:
                    text_fragment = html_lib.escape(run.text)
                    if run.bold and run.italic:
                        text_fragment = f"<strong><em>{text_fragment}</em></strong>"
                    elif run.bold:
                        text_fragment = f"<strong>{text_fragment}</strong>"
                    elif run.italic:
                        text_fragment = f"<em>{text_fragment}</em>"
                    parts_html.append(text_fragment)
            content = "".join(parts_html) if parts_html else html_lib.escape(text)
            html_parts.append(f'<p style="margin:0.4em 0;line-height:1.7">{content}</p>')

    for table in doc.tables:
        html_parts.append('<table style="border-collapse:collapse;width:100%;margin:1em 0">')
        for row_index, row in enumerate(table.rows):
            html_parts.append("<tr>")
            tag = "th" if row_index == 0 else "td"
            style_attr = (
                "padding:6px 10px;border:1px solid #ccc;font-weight:bold;background:#f5f5f5"
                if row_index == 0
                else "padding:6px 10px;border:1px solid #ccc"
            )
            for cell in row.cells:
                html_parts.append(f'<{tag} style="{style_attr}">{html_lib.escape(cell.text)}</{tag}>')
            html_parts.append("</tr>")
        html_parts.append("</table>")

    html_fragment = "\n".join(html_parts) if html_parts else "<p><em>Document vide ou sans contenu lisible.</em></p>"
    return {"html": html_fragment, "messages": []}


@router.get("/all/user")
async def list_all_user_files(authorization: Optional[str] = Header(None)):
    """Liste tous les fichiers de l'utilisateur, toutes conversations confondues."""
    user_id = get_user_id_from_header(authorization)
    if user_id == "anonymous":
        raise HTTPException(status_code=403, detail="Authentification requise.")

    db = SessionLocal()
    try:
        conversations = (
            db.query(Conversation)
            .filter(Conversation.user_id == user_id)
            .order_by(Conversation.updated_at.desc())
            .all()
        )
        conv_titles = {conv.id: conv.title for conv in conversations}
        conv_ids = list(conv_titles.keys())

        stored_rows = (
            db.query(ConversationFile)
            .filter(ConversationFile.user_id == user_id)
            .order_by(ConversationFile.created_at.desc())
            .limit(200)
            .all()
        )

        attachment_rows: List[Message] = []
        if conv_ids:
            attachment_rows = (
                db.query(Message)
                .filter(
                    Message.conversation_id.in_(conv_ids),
                    Message.attachment_json.isnot(None),
                )
                .order_by(Message.timestamp.desc(), Message.id.desc())
                .limit(200)
                .all()
            )

        result = _merge_file_entries(
            [_serialize_conversation_file(row, conv_titles) for row in stored_rows]
            + [_serialize_message_attachment(row, conv_titles) for row in attachment_rows]
        )
        return {"files": result, "count": len(result)}
    finally:
        db.close()


@router.get("/proxy")
async def proxy_file(url: str, download: int = 0, name: Optional[str] = None):
    """
    Proxy un fichier cote serveur pour eviter les erreurs CORS.
    ?download=1 force le telechargement.
    """
    import urllib.parse

    import httpx
    from fastapi.responses import Response

    filename = urllib.parse.unquote(name or url.split("/")[-1].split("?")[0])

    try:
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as http:
            resp = await http.get(url)
            if resp.status_code == 404:
                raise HTTPException(status_code=404, detail="Fichier introuvable (404).")
            resp.raise_for_status()
            content = resp.content
            content_type = resp.headers.get("content-type", "application/octet-stream")
    except HTTPException:
        raise
    except Exception as exc:
        logger.error(f"[proxy] Erreur telechargement {url}: {exc}")
        raise HTTPException(status_code=502, detail=f"Proxy error: {str(exc)}")

    disposition = "attachment" if download else "inline"
    return Response(
        content=content,
        media_type=content_type,
        headers={
            "Content-Disposition": f'{disposition}; filename="{filename}"',
            "Access-Control-Allow-Origin": "*",
            "Cache-Control": "public, max-age=3600",
        },
    )


@router.get("/render-doc")
async def render_doc_as_html(url: str, authorization: Optional[str] = Header(None)):
    try:
        file_bytes = await _load_file_bytes(url=url)
        return _render_docx_bytes_as_html(file_bytes)
    except Exception as exc:
        logger.error(f"render-doc error: {exc}", exc_info=True)
        raise HTTPException(status_code=502, detail=f"Erreur conversion document : {str(exc)}")


@router.post("/render-doc-inline")
async def render_doc_inline(payload: InlineFilePayload, authorization: Optional[str] = Header(None)):
    try:
        file_bytes = await _load_file_bytes(data_url=payload.data_url)
        return _render_docx_bytes_as_html(file_bytes)
    except Exception as exc:
        logger.error(f"render-doc-inline error: {exc}", exc_info=True)
        raise HTTPException(status_code=502, detail=f"Erreur conversion document inline : {str(exc)}")


@router.get("/render-sheet")
async def render_sheet_as_json(url: str, authorization: Optional[str] = Header(None)):
    import json

    from agents.workers.spreadsheet_worker import read_excel_document_as_json

    try:
        json_str = await read_excel_document_as_json.ainvoke({"file_path_or_url": url})
        if json_str.startswith("Erreur"):
            raise HTTPException(status_code=500, detail=json_str)
        return json.loads(json_str)
    except Exception as exc:
        logger.error(f"render-sheet error: {exc}", exc_info=True)
        raise HTTPException(status_code=502, detail=f"Erreur conversion tableur : {str(exc)}")


@router.post("/render-sheet-inline")
async def render_sheet_inline(payload: InlineFilePayload, authorization: Optional[str] = Header(None)):
    import json

    from agents.workers.spreadsheet_worker import read_excel_document_as_json

    try:
        json_str = await read_excel_document_as_json.ainvoke({"file_path_or_url": payload.data_url})
        if json_str.startswith("Erreur"):
            raise HTTPException(status_code=500, detail=json_str)
        return json.loads(json_str)
    except Exception as exc:
        logger.error(f"render-sheet-inline error: {exc}", exc_info=True)
        raise HTTPException(status_code=502, detail=f"Erreur conversion tableur inline : {str(exc)}")


@router.get("/structure")
async def get_file_structure(url: str, authorization: Optional[str] = Header(None)):
    from docx import Document
    from openpyxl import load_workbook

    try:
        file_bytes = io.BytesIO(await _load_file_bytes(url=url))

        if url.endswith(".docx") or "document" in url:
            doc = Document(file_bytes)
            sections = []
            for index, para in enumerate(doc.paragraphs):
                if para.style.name.startswith("Heading") or para.style.name == "Title":
                    sections.append({"id": f"p_{index}", "label": para.text[:50]})
            for index, table in enumerate(doc.tables):
                sections.append({"id": f"t_{index}", "label": f"Tableau {index + 1} (Lignes: {len(table.rows)})"})
            return {"type": "document", "sections": sections}

        if url.endswith(".xlsx") or "sheet" in url:
            workbook = load_workbook(file_bytes, data_only=True)
            return {"type": "sheet", "sheets": workbook.sheetnames}

        return {"type": "unknown"}
    except Exception as exc:
        return {"error": str(exc)}


@router.post("/structure-inline")
async def get_inline_file_structure(payload: InlineFilePayload, authorization: Optional[str] = Header(None)):
    from docx import Document
    from openpyxl import load_workbook

    try:
        file_bytes = io.BytesIO(await _load_file_bytes(data_url=payload.data_url))
        file_name = (payload.name or "").lower()

        if file_name.endswith(".docx"):
            doc = Document(file_bytes)
            sections = []
            for index, para in enumerate(doc.paragraphs):
                if para.style.name.startswith("Heading") or para.style.name == "Title":
                    sections.append({"id": f"p_{index}", "label": para.text[:50]})
            for index, table in enumerate(doc.tables):
                sections.append({"id": f"t_{index}", "label": f"Tableau {index + 1} (Lignes: {len(table.rows)})"})
            return {"type": "document", "sections": sections or ["Document entier"]}

        if file_name.endswith(".xlsx") or file_name.endswith(".xls"):
            workbook = load_workbook(file_bytes, data_only=True)
            return {"type": "sheet", "sheets": workbook.sheetnames}

        return {"type": "unknown"}
    except Exception as exc:
        return {"error": str(exc)}


@router.get("/{conversation_id}")
async def list_files(conversation_id: str, authorization: Optional[str] = Header(None)):
    """Liste tous les fichiers d'une conversation."""
    user_id = get_user_id_from_header(authorization)

    db = SessionLocal()
    try:
        conversation = db.query(Conversation).filter(Conversation.id == conversation_id).first()
        if not conversation:
            raise HTTPException(status_code=404, detail="Conversation introuvable.")
        if user_id != "anonymous" and conversation.user_id != user_id:
            raise HTTPException(status_code=403, detail="Acces refuse.")

        files_query = db.query(ConversationFile).filter(ConversationFile.conversation_id == conversation_id)
        if user_id != "anonymous":
            files_query = files_query.filter(ConversationFile.user_id == user_id)

        stored_rows = files_query.order_by(ConversationFile.created_at.desc()).all()
        attachment_rows = (
            db.query(Message)
            .filter(
                Message.conversation_id == conversation_id,
                Message.attachment_json.isnot(None),
            )
            .order_by(Message.timestamp.desc(), Message.id.desc())
            .all()
        )

        conv_titles = {conversation_id: conversation.title or ""}
        result = _merge_file_entries(
            [_serialize_conversation_file(row, conv_titles) for row in stored_rows]
            + [_serialize_message_attachment(row, conv_titles) for row in attachment_rows]
        )
        return {"conversation_id": conversation_id, "files": result, "count": len(result)}
    finally:
        db.close()


@router.delete("/{conversation_id}/{file_name}")
async def remove_file(conversation_id: str, file_name: str, authorization: Optional[str] = Header(None)):
    """Supprime un fichier d'une conversation."""
    user_id = get_user_id_from_header(authorization)
    if user_id == "anonymous":
        raise HTTPException(status_code=403, detail="Authentification requise.")

    storage_path = f"users/{user_id}/conversations/{conversation_id}/{file_name}"
    ok = storage.delete_file(storage_path)

    if not ok:
        raise HTTPException(status_code=404, detail="Fichier introuvable.")

    return {"status": "deleted", "file_name": file_name}
