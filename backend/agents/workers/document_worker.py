"""
Worker Document — FLARE AI.
Gère : Création, édition et formatage de documents Microsoft Word (.docx).
"""
import asyncio
import logging
import json
import uuid
import io
import base64
import operator
import os
from typing import TypedDict, Annotated, Sequence, Literal, Optional, List

from langchain_core.messages import BaseMessage, HumanMessage, AIMessage, SystemMessage
from langchain_core.tools import tool
from langgraph.graph import StateGraph, END
from langgraph.prebuilt import ToolNode
from langchain_core.runnables import RunnableConfig

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from core.llm_factory import get_llm
from core.memory import SessionMemory
from core.config import settings

logger = logging.getLogger(__name__)

from core.context import (
    current_user_id as _current_user_id,
    current_session_id as _current_session_id,
    current_request_id as _current_request_id,
    generated_images as _generated_images,
    GLOBAL_IMAGE_REGISTRY as _GLOBAL_IMAGE_REGISTRY,
    current_inline_file as _current_inline_file,
)


def _get_inline_file_payload() -> Optional[dict]:
    payload = _current_inline_file.get()
    return payload if isinstance(payload, dict) else None


def _extract_filename_from_source(file_path_or_url: str, fallback: str) -> str:
    if file_path_or_url.startswith("inline://"):
        payload = _get_inline_file_payload()
        return str(payload.get("name") or fallback) if payload else fallback
    if file_path_or_url.startswith("data:"):
        payload = _get_inline_file_payload()
        return str(payload.get("name") or fallback) if payload else fallback
    if "/" in file_path_or_url:
        return file_path_or_url.split("/")[-1].split("?")[0] or fallback
    return file_path_or_url or fallback


async def _load_docx_bytes(file_path_or_url: str) -> bytes:
    if file_path_or_url.startswith("inline://"):
        payload = _get_inline_file_payload()
        if not payload or not payload.get("content"):
            raise ValueError("Aucun fichier inline n'est disponible dans ce contexte.")
        return base64.b64decode(payload["content"])

    if file_path_or_url.startswith("data:"):
        try:
            _, raw_b64 = file_path_or_url.split(",", 1)
            return base64.b64decode(raw_b64)
        except Exception as exc:
            raise ValueError("Le document inline est invalide.") from exc

    if file_path_or_url.startswith("http://") or file_path_or_url.startswith("https://"):
        import httpx
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
            response = await client.get(file_path_or_url)
            response.raise_for_status()
            return response.content

    with open(file_path_or_url, "rb") as f:
        return f.read()

@tool
async def generate_word_document(filename: str, elements_json: str, config: RunnableConfig) -> str:
    """Crée ou met à jour un document Microsoft Word (.docx) avec formatage avancé.

    Args:
        filename: Nom du fichier (doit se terminer par .docx). Exemple : "Rapport.docx"
        elements_json: Chaîne JSON représentant une liste de blocs.
    """
    configurable = config.get("configurable", {}) if config else {}     
    user_id = configurable.get("user_id") or _current_user_id.get() or "anonymous"
    session_id = configurable.get("session_id") or _current_session_id.get() or "default"
    req_id = configurable.get("request_id") or _current_request_id.get()

    if req_id and req_id in _GLOBAL_IMAGE_REGISTRY:
        current_files = _GLOBAL_IMAGE_REGISTRY[req_id]
    else:
        current_files = _generated_images.get() or []

    try:
        elements = json.loads(elements_json)
    except json.JSONDecodeError as e:
        return f"Erreur : Le format JSON fourni est invalide. Détail : {str(e)}"

    if not filename.endswith(".docx"):
        filename += ".docx"

    doc = Document()

    def hex_to_rgb(hex_str):
        hex_str = hex_str.lstrip('#')
        if len(hex_str) != 6:
            return RGBColor(0, 0, 0)
        return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

    for el in elements:
        t = el.get("type")
        text = el.get("text", "")
        align_str = el.get("align", "left").lower()

        alignment = WD_ALIGN_PARAGRAPH.LEFT
        if align_str == "center":
            alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align_str == "right":
            alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align_str == "justify":
            alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if t == "title":
            p = doc.add_heading(text, level=0)
            p.alignment = alignment
            if el.get("color"):
                for run in p.runs:
                    run.font.color.rgb = hex_to_rgb(el["color"])        

        elif t == "heading":
            level = min(max(el.get("level", 1), 1), 9)
            p = doc.add_heading(text, level=level)
            p.alignment = alignment
            if el.get("color"):
                for run in p.runs:
                    run.font.color.rgb = hex_to_rgb(el["color"])        

        elif t == "paragraph":
            p = doc.add_paragraph()
            p.alignment = alignment
            run = p.add_run(text)
            run.bold = el.get("bold", False)
            run.italic = el.get("italic", False)
            if el.get("color"):
                run.font.color.rgb = hex_to_rgb(el["color"])
            
            if el.get("review_notes"):
                run = p.add_run(f" [NOTE IA : {el.get('review_notes')}]")
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.italic = True

        elif t == "list":
            style = 'List Bullet' if el.get("style") == "bullet" else 'List Number'
            items = el.get("items", [])
            for item in items:
                doc.add_paragraph(str(item), style=style)

        elif t == "table":
            headers = el.get("headers", [])
            rows = el.get("rows", [])
            if headers or rows:
                num_cols = len(headers) if headers else len(rows[0]) if rows else 1
                table = doc.add_table(rows=1 if headers else 0, cols=num_cols)
                table.style = 'Table Grid'

                if headers:
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        if i < len(hdr_cells):
                            hdr_cells[i].text = str(header)
                            for paragraph in hdr_cells[i].paragraphs:   
                                for run in paragraph.runs:
                                    run.font.bold = True

                for row_data in rows:
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row_data):
                        if i < len(row_cells):
                            row_cells[i].text = str(val)
        
        elif t == "image":
            base64_data = el.get("data_base64", el.get("base64_data"))
            width_inches = el.get("width_inches", 5.0)
            if base64_data:
                import base64
                try:
                    image_data = base64.b64decode(base64_data)
                    doc.add_picture(io.BytesIO(image_data), width=Inches(width_inches))
                except Exception as e:
                    logger.error(f"Erreur lors de l'insertion de l'image base64: {e}")

        elif t == "cover_page":
            # Page de couverture stylée
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph()
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.add_run(el.get("title", "Document"))
            title_run.font.size = Pt(32)
            title_run.bold = True
            color_str = el.get("color", "1B2A4A")
            title_run.font.color.rgb = hex_to_rgb(color_str)
            if el.get("subtitle"):
                sub_p = doc.add_paragraph()
                sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub_run = sub_p.add_run(el["subtitle"])
                sub_run.font.size = Pt(18)
                sub_run.font.color.rgb = hex_to_rgb(el.get("accent_color", "4A6572"))
            doc.add_paragraph()
            doc.add_paragraph()
            if el.get("author") or el.get("date"):
                meta_p = doc.add_paragraph()
                meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                meta_text = " — ".join(filter(None, [el.get("author"), el.get("date")]))
                meta_run = meta_p.add_run(meta_text)
                meta_run.font.size = Pt(12)
                meta_run.font.color.rgb = hex_to_rgb("888888")
            doc.add_page_break()

        elif t == "callout":
            # Encadré coloré (tip, warning, example, note)
            style_map = {
                "tip":     ("2563EB", "EFF6FF", "💡 "),
                "warning": ("D97706", "FFFBEB", "⚠️ "),
                "example": ("059669", "ECFDF5", "📌 "),
                "note":    ("6B7280", "F9FAFB", "📝 "),
            }
            border_color, bg_hint, icon = style_map.get(el.get("style", "tip"), style_map["tip"])
            callout_title = el.get("title", "")
            callout_text = el.get("text", "")
            p = doc.add_paragraph()
            if callout_title:
                r = p.add_run(f"{icon}{callout_title}")
                r.bold = True
                r.font.color.rgb = hex_to_rgb(border_color)
            if callout_text:
                if callout_title:
                    p2 = doc.add_paragraph()
                else:
                    p2 = p
                r2 = p2.add_run(callout_text)
                r2.font.color.rgb = hex_to_rgb("374151")

        elif t == "page_break":
            doc.add_page_break()

        elif t == "divider":
            p = doc.add_paragraph("─" * 60)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif t == "toc":
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar)
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = r'TOC \o "1-3" \h \z \u'
            run._r.append(instrText)
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'separate')
            run._r.append(fldChar)
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_bytes = doc_io.getvalue()

    try:
        file_uuid = str(uuid.uuid4())[:8]
        safe_filename = filename.replace(" ", "_")
        storage_path = f"users/{user_id}/conversations/{session_id}/doc_{file_uuid}_{safe_filename}"

        from core.firebase_client import firebase_storage as storage    
        public_url = storage.upload_file(
            bucket_name=settings.NEXT_PUBLIC_FIREBASE_STORAGE_BUCKET,   
            path=storage_path, 
            file_bytes=doc_bytes,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        if public_url:
            try:
                memory = SessionMemory(session_id=session_id, user_id=user_id)
                memory.save_file_record(file_name=safe_filename, file_url=public_url,
                                        file_type="document", mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        file_size=len(doc_bytes))       
            except Exception as e:
                logger.error(f"Erreur SQL document_worker: {e}")        

        import base64
        b64_doc = base64.b64encode(doc_bytes).decode('utf-8')
        doc_obj = {
            "name": safe_filename,
            "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "url": public_url,
            "data": b64_doc if not public_url else None
        }
        current_files.append(doc_obj)
        if req_id:
            _GLOBAL_IMAGE_REGISTRY[req_id] = current_files
        _generated_images.set(current_files)

        # ── Log coût document dans UsageLedger ─────────────────────────
        try:
            from core.database import SessionLocal as _SL, UsageLedger as _UL
            from core.config import MEDIA_PRICING
            import uuid as _uuid
            _db = _SL()
            _db.add(_UL(
                id=str(_uuid.uuid4()),
                user_id=user_id,
                model_name="document-worker",
                action_kind="doc_gen",
                prompt_tokens=0, candidate_tokens=0, total_tokens=0,
                cost_usd=MEDIA_PRICING.get("doc_gen", 0.005),
                usage_metadata={"filename": safe_filename, "elements": len(elements)},
            ))
            _db.commit()
            _db.close()
        except Exception as _e:
            logger.warning(f"[generate_word_document] Erreur log UsageLedger: {_e}")

        if public_url:
            result_message = f"Document '{safe_filename}' généré avec succès. URL: {public_url}"
        else:
            result_message = f"Document '{safe_filename}' généré. Il a été ajouté aux fichiers de la session."

        return result_message

    except Exception as e:
        logger.error(f"Erreur persistence generate_word_document: {e}")
        return f"Le document a été généré mais une erreur est survenue lors de l'enregistrement: {e}"


@tool
async def generate_chart_image(chart_type: str, data: List[float], labels: List[str], title: str) -> str:
    """Génère une image de graphique à partir de données structurées et la retourne en base64.

    Args:
        chart_type: Type de graphique ('bar', 'pie', 'line').
        data: Liste des valeurs numériques pour le graphique.
        labels: Liste des étiquettes pour chaque valeur.
        title: Titre du graphique.

    Returns:
        Une chaîne base64 représentant l'image PNG du graphique.
    """
    try:
        fig, ax = plt.subplots()

        if chart_type == 'bar':
            ax.bar(labels, data)
        elif chart_type == 'pie':
            ax.pie(data, labels=labels, autopct='%1.1f%%')
        elif chart_type == 'line':
            ax.plot(labels, data)
        else:
            return "Erreur: Type de graphique non supporté. Utilisez 'bar', 'pie', ou 'line'."

        ax.set_title(title)
        plt.tight_layout()

        img_io = io.BytesIO()
        plt.savefig(img_io, format='png')
        img_io.seek(0)
        plt.close(fig)

        import base64
        base64_string = base64.b64encode(img_io.getvalue()).decode('utf-8')
        return base64_string
    except Exception as e:
        logger.error(f"Erreur lors de la génération du graphique: {e}")
        return f"Erreur interne lors de la création du graphique: {e}"


@tool
async def read_word_document_as_json(file_path_or_url: str) -> str:     
    """Lit un document Word (.docx) et retourne sa structure en JSON, prêt à être modifié et utilisé par update_word_document.

    Args:
        file_path_or_url: Le chemin local ou l'URL du fichier .docx à lire.    

    Returns:
        Une chaîne JSON représentant les éléments du document.      
    """
    try:
        document = Document(io.BytesIO(await _load_docx_bytes(file_path_or_url)))
        elements = []
        for para in document.paragraphs:
            if not para.text.strip():
                continue

            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split(' ')[-1])
                elements.append({
                    "type": "heading",
                    "level": level,
                    "text": para.text,
                })
            elif para.style.name == 'Title':
                 elements.append({
                    "type": "title",
                    "text": para.text,
                })
            elif para.style.name.startswith('List'):
                if elements and elements[-1]["type"] == "list" and elements[-1].get("style_name") == para.style.name:
                    elements[-1]["items"].append(para.text)
                else:
                    elements.append({
                        "type": "list",
                        "items": [para.text],
                        "style": "bullet" if "Bullet" in para.style.name else "number",
                        "style_name": para.style.name   
                    })
            else:
                elements.append({
                    "type": "paragraph",
                    "text": para.text,
                })

        for el in elements:
            if "style_name" in el:
                del el["style_name"]

        for table in document.tables:
            headers = [cell.text for cell in table.rows[0].cells]       
            rows = []
            for row in table.rows[1:]:
                rows.append([cell.text for cell in row.cells])
            elements.append({
                "type": "table",
                "headers": headers,
                "rows": rows
            })

        return json.dumps(elements, indent=2, ensure_ascii=False)       

    except Exception as e:
        logger.error(f"Erreur lors de la lecture JSON du document Word: {e}")
        return f"Erreur de lecture du document en JSON: {e}"


@tool
async def read_word_document(file_path_or_url: str) -> str:
    """Lit le contenu textuel d'un document Microsoft Word (.docx) à partir d'un chemin local ou d'une URL.

    Args:
        file_path_or_url: Le chemin local ou l'URL du fichier .docx à lire.    

    Returns:
        Le contenu textuel extrait du document.
    """
    try:
        document = Document(io.BytesIO(await _load_docx_bytes(file_path_or_url)))

        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)

        return "\n".join(full_text)



    except Exception as e:
        logger.error(f"Erreur lors de la lecture du document Word: {e}")
        return f"Erreur de lecture du document: {e}"

@tool
async def update_word_document(file_path_or_url: str, elements_json: str, config: RunnableConfig) -> str:
    """Met à jour un document Word (.docx) existant.

    Args:
        file_path_or_url: L'URL ou le chemin du document à modifier.
        elements_json: La nouvelle structure JSON du document (obtenue après modification du JSON retourné par read_word_document_as_json).   
    """
    filename = _extract_filename_from_source(file_path_or_url, "Document_mis_a_jour.docx")
    if filename.startswith("doc_"):
        parts = filename.split("_", 2)
        if len(parts) >= 3:
            filename = parts[-1]

    result = await generate_word_document.ainvoke({"filename": filename, "elements_json": elements_json}, config=config)

    if file_path_or_url.startswith("http"):
        try:
            from core.firebase_client import firebase_storage
            bucket_url_prefix = f"https://storage.googleapis.com/{settings.NEXT_PUBLIC_FIREBASE_STORAGE_BUCKET}/"
            if file_path_or_url.startswith(bucket_url_prefix):
                path_to_delete = file_path_or_url.replace(bucket_url_prefix, "")
                firebase_storage.delete_file(path_to_delete)
        except Exception as e:
            logger.error(f"Failed to delete old cloud file: {e}")       
            pass

    return f"Mise à jour réussie. {result}"

@tool
async def delete_word_document(file_path_or_url: str) -> str:
    """Supprime un document Word (.docx) existant.

    Args:
        file_path_or_url: L'URL ou le chemin local du document à supprimer.    
    """
    if file_path_or_url.startswith("inline://") or file_path_or_url.startswith("data:"):
        return "Document inline supprimé du contexte courant."
    if file_path_or_url.startswith("http"):
        try:
            from core.firebase_client import firebase_storage
            bucket_url_prefix = f"https://storage.googleapis.com/{settings.NEXT_PUBLIC_FIREBASE_STORAGE_BUCKET}/"
            if file_path_or_url.startswith(bucket_url_prefix):
                path_to_delete = file_path_or_url.replace(bucket_url_prefix, "")
                success = firebase_storage.delete_file(path_to_delete)  
                if success:
                    return f"Le document a été supprimé avec succès du stockage cloud."
                else:
                    return "Échec de la suppression du document cloud."
            else:
                return "L'URL fournie ne correspond pas au stockage de l'application. Suppression ignorée."
        except Exception as e:
            logger.error(f"Erreur lors de la suppression cloud: {e}")   
            return f"Erreur lors de la suppression cloud: {e}"
    else:
        try:
            if os.path.exists(file_path_or_url):
                os.remove(file_path_or_url)
                return f"Document local {file_path_or_url} supprimé avec succès."
            else:
                return f"Le fichier local {file_path_or_url} n'existe pas."
        except Exception as e:
            return f"Erreur lors de la suppression locale: {e}"

DOCUMENT_TOOLS = [
    generate_word_document,
    generate_chart_image,
    read_word_document_as_json,
    read_word_document,
    update_word_document,
    delete_word_document,
]

DOCUMENT_SYSTEM_PROMPT = """Tu es l'Agent Créateur de Documents de FLARE AI. Tu produis des documents Word (.docx) PROFESSIONNELS, DENSES et BIEN STRUCTURÉS.

━━━ CRUD ━━━
1. **CREATE**: Conçois une structure JSON riche et utilise `generate_word_document`.
2. **READ**: Utilise `read_word_document` (texte brut) ou `read_word_document_as_json` (structure).
3. **UPDATE** — C'est le cas le plus important :
   - Si le message contient `[MODIFICATION FICHIER]`, extrais l'URL du fichier depuis ce bloc.
   - Appelle `read_word_document_as_json` avec cette URL pour lire la structure actuelle.
   - Si `[SELECTION]...[/SELECTION]` est présent : modifie UNIQUEMENT le passage sélectionné, retrouve-le dans le JSON et applique la modification chirurgicalement.
   - Sinon : modifie les sections nécessaires sans tout réécrire.
   - Appelle `update_word_document` avec l'URL originale + JSON modifié.
4. **DELETE**: Utilise `delete_word_document` avec l'URL.

━━━ TYPES D'ÉLÉMENTS DISPONIBLES ━━━
```json
{"type": "cover_page", "title": "Titre", "subtitle": "Sous-titre", "author": "Auteur", "date": "2024"}
{"type": "toc"}
{"type": "title", "text": "...", "color": "1B2A4A", "align": "center"}
{"type": "heading", "level": 1, "text": "...", "color": "1B2A4A"}
{"type": "paragraph", "text": "...", "bold": false, "italic": false, "color": "333333"}
{"type": "list", "style": "bullet", "items": ["Point 1", "Point 2"]}
{"type": "list", "style": "number", "items": ["Étape 1", "Étape 2"]}
{"type": "table", "headers": ["Col A", "Col B"], "rows": [["val1", "val2"]]}
{"type": "callout", "style": "tip", "title": "Bon à savoir", "text": "..."}
{"type": "callout", "style": "warning", "title": "Attention", "text": "..."}
{"type": "callout", "style": "example", "title": "Exemple", "text": "..."}
{"type": "page_break"}
{"type": "image", "data_base64": "...", "width_inches": 6.0}
```

━━━ EXIGENCES QUALITÉ ━━━
- **Densité**: Chaque paragraphe doit contenir 3-6 phrases complètes. Pas de paragraphes d'une ligne.
- **Structure systématique**: Tout document doit avoir au minimum: cover_page → toc → sections numérotées avec headings → conclusion.
- **Contenu réel**: Rédige du VRAI contenu, pas des placeholders. Si c'est un cours, rédige vraiment le cours.
- **Cohérence**: Utilise le même thème de couleur dans tout le document.
- **Callouts**: Utilise-les pour mettre en valeur points clés, avertissements, exemples pratiques.
- **Longueur**: Un bon document = 8-20 éléments minimum (title + toc + 3-5 sections avec plusieurs paragraphes chacune).

━━━ THÈMES ━━━
- **Corporate** : bleu marine `1B2A4A`, gris `6B7280`, accents `2563EB`
- **Créatif** : orange `F97316`, teal `0D9488`, marine `1E293B`
- **Académique** : noir `111827`, bordeaux `9B1C1C`, gris `374151`

━━━ FORMAT DE RÉPONSE ━━━
Réponds toujours brièvement (1-2 phrases max). Propose 2 suggestions avec `[SUGGESTION: ...]`.
"""

class DocumentWorker:
    """Worker spécialisé dans la création de documents — graphe LangGraph autonome."""

    def __init__(self, model_override: str = None):
        self.tools = DOCUMENT_TOOLS
        self.llm = get_llm(
            temperature=0.2,
            model_override=model_override or "gemini-2.5-flash",    
        ).bind_tools(self.tools)
        self.tool_node = ToolNode(self.tools)
        self.graph = self._build_graph()
        logger.info(f"[DocumentWorker] Initialisé avec {len(self.tools)} outils")

    def _build_graph(self):
        graph = StateGraph(TypedDict("DocumentState", {"messages": Annotated[Sequence[BaseMessage], operator.add]}))
        graph.add_node("agent", self._call_model)
        graph.add_node("tools", self.tool_node)
        graph.set_entry_point("agent")
        graph.add_conditional_edges("agent", self._should_continue, {"continue": "tools", "end": END})
        graph.add_edge("tools", "agent")
        return graph.compile()

    def _should_continue(self, state) -> Literal["continue", "end"]:    
        last = state["messages"][-1]
        if hasattr(last, "tool_calls") and last.tool_calls:
            return "continue"
        return "end"

    async def _call_model(self, state, config: RunnableConfig = None) -> dict:
        messages = state["messages"]

        if not any(isinstance(m, SystemMessage) for m in messages):     
            messages = [SystemMessage(content=DOCUMENT_SYSTEM_PROMPT)] + list(messages)

        for attempt in range(3):
            try:
                return {"messages": [await self.llm.ainvoke(messages)]} 
            except Exception as e:
                if attempt < 2 and any(k in str(e).lower() for k in ["429", "500", "503"]):
                    await asyncio.sleep(2 ** (attempt + 1))
                    continue
                raise

    async def run(self, task: str, config: dict = None) -> str:
        prompt_content = f"""[Instructions]
{DOCUMENT_SYSTEM_PROMPT}
[Fin instructions]

{task}"""








        messages = [HumanMessage(content=prompt_content)]
        result = await self.graph.ainvoke({"messages": messages}, config=config or {})
        last_msg = result["messages"][-1]
        content = getattr(last_msg, "content", "")
        if isinstance(content, list):
            content = " ".join(part.get("text", "") if isinstance(part, dict) else str(part) for part in content)
        return content or "Tâche de document terminée."






